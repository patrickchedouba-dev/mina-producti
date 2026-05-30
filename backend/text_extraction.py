"""
Extraction de texte depuis différents formats de documents.
Supporte PDF, TXT, DOCX.
Pour les PDFs scannés (images), utilise Google Document AI OCR.
"""

import io
import logging
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """
    Résultat de l'extraction de texte.
    """
    text: str
    page_count: Optional[int] = None
    success: bool = True
    error_message: Optional[str] = None
    used_ocr: bool = False


class TextExtractor:
    """
    Extracteur de texte multi-format avec support Document AI OCR.
    
    Document AI traite les PDFs côté serveur Google, donc pas de
    problème de mémoire locale contrairement à Vision + pdf2image.
    """
    
    def __init__(
        self,
        use_ocr: bool = True,
        project_id: str = None,
        location: str = "us",
        processor_id: str = None
    ):
        """
        Initialise l'extracteur.
        
        Args:
            use_ocr: Activer l'OCR pour les PDFs sans texte
            project_id: ID du projet GCP
            location: Région du processeur Document AI (us ou eu)
            processor_id: ID du processeur Document AI (OCR)
        """
        self.use_ocr = use_ocr
        self.project_id = project_id
        self.location = location
        self.processor_id = processor_id
        self._docai_client = None
        
        logger.info(
            f"TextExtractor initialisé "
            f"(OCR={'Document AI' if use_ocr else 'disabled'})"
        )
    
    def _init_document_ai(self):
        """Initialise le client Document AI si nécessaire."""
        if self._docai_client is not None:
            return True
            
        if not self.use_ocr:
            return False
        
        try:
            from google.cloud import documentai_v1 as documentai
            
            # Obtenir le project_id depuis l'environnement si pas fourni
            if not self.project_id:
                import google.auth
                _, self.project_id = google.auth.default()
            
            self._docai_client = documentai.DocumentProcessorServiceClient()
            
            # Créer ou obtenir un processeur OCR
            if not self.processor_id:
                self.processor_id = self._get_or_create_ocr_processor()
            
            if self.processor_id:
                logger.info(f"Document AI initialisé: processor={self.processor_id}")
                return True
            else:
                logger.warning("Impossible de créer un processeur Document AI")
                self.use_ocr = False
                return False
                
        except ImportError:
            logger.warning("google-cloud-documentai non installé, OCR désactivé")
            self.use_ocr = False
            return False
        except Exception as e:
            logger.warning(f"Erreur init Document AI: {e}")
            self.use_ocr = False
            return False
    
    def _get_or_create_ocr_processor(self) -> Optional[str]:
        """Obtient ou crée un processeur OCR."""
        from google.cloud import documentai_v1 as documentai
        
        parent = f"projects/{self.project_id}/locations/{self.location}"
        
        try:
            # Lister les processeurs existants
            processors = list(self._docai_client.list_processors(parent=parent))
            
            # Chercher un processeur OCR existant
            for processor in processors:
                if "OCR" in processor.type_.upper() or "ocr" in processor.display_name.lower():
                    logger.info(f"Processeur OCR trouvé: {processor.name}")
                    return processor.name
            
            # Créer un nouveau processeur OCR
            logger.info("Création d'un nouveau processeur OCR...")
            
            processor = documentai.Processor(
                display_name="mina-ocr-processor",
                type_="OCR_PROCESSOR",
            )
            
            result = self._docai_client.create_processor(
                parent=parent,
                processor=processor,
            )
            
            logger.info(f"Processeur OCR créé: {result.name}")
            return result.name
            
        except Exception as e:
            logger.error(f"Erreur get/create processeur: {e}")
            return None
    

    def extract_from_uri(self, gcs_uri: str, filename: str) -> ExtractionResult:
        """Extrait le texte d'un PDF via son URI GCS — sans téléchargement local."""
        if not self._init_document_ai():
            return ExtractionResult(text='', success=False, error_message='Document AI non disponible')
        try:
            from google.cloud import documentai_v1 as documentai
            from google.cloud import storage
            import json, time
            output_prefix = f"_docai_output/{filename.replace('/','_').replace(' ','_')}/"
            output_uri = f"gs://mina-pdfs/{output_prefix}"
            request = documentai.BatchProcessRequest(
                name=self.processor_id,
                input_documents=documentai.BatchDocumentsInputConfig(
                    gcs_documents=documentai.GcsDocuments(
                        documents=[documentai.GcsDocument(gcs_uri=gcs_uri, mime_type='application/pdf')]
                    )
                ),
                document_output_config=documentai.DocumentOutputConfig(
                    gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(gcs_uri=output_uri)
                )
            )
            print(f'[DEBUG] Batch lancé: {gcs_uri}')
            operation = self._docai_client.batch_process_documents(request=request)
            operation.result(timeout=300)
            storage_client = storage.Client()
            bucket = storage_client.bucket('mina-pdfs')
            full_text = []
            for blob in bucket.list_blobs(prefix=output_prefix):
                if blob.name.endswith('.json'):
                    data = json.loads(blob.download_as_text())
                    full_text.append(data.get('text', ''))
                    blob.delete()
            text = self._clean_text('\n'.join(full_text))
            logger.info(f'Extraction [GCS batch]: {filename} ({len(text)} chars)')
            return ExtractionResult(text=text, success=True)
        except Exception as e:
            logger.error(f'Erreur batch GCS {filename}: {e}')
            return ExtractionResult(text='', success=False, error_message=str(e))


    def extract_from_uri(self, gcs_uri: str, filename: str) -> ExtractionResult:
        """Extrait le texte d'un PDF via son URI GCS — sans téléchargement local."""
        if not self._init_document_ai():
            return ExtractionResult(text='', success=False, error_message='Document AI non disponible')
        try:
            from google.cloud import documentai_v1 as documentai
            from google.cloud import storage
            import json, time
            output_prefix = f"_docai_output/{filename.replace('/','_').replace(' ','_')}/"
            output_uri = f"gs://mina-pdfs/{output_prefix}"
            request = documentai.BatchProcessRequest(
                name=self.processor_id,
                input_documents=documentai.BatchDocumentsInputConfig(
                    gcs_documents=documentai.GcsDocuments(
                        documents=[documentai.GcsDocument(gcs_uri=gcs_uri, mime_type='application/pdf')]
                    )
                ),
                document_output_config=documentai.DocumentOutputConfig(
                    gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(gcs_uri=output_uri)
                )
            )
            print(f'[DEBUG] Batch lancé: {gcs_uri}')
            operation = self._docai_client.batch_process_documents(request=request)
            operation.result(timeout=300)
            storage_client = storage.Client()
            bucket = storage_client.bucket('mina-pdfs')
            full_text = []
            for blob in bucket.list_blobs(prefix=output_prefix):
                if blob.name.endswith('.json'):
                    data = json.loads(blob.download_as_text())
                    full_text.append(data.get('text', ''))
                    blob.delete()
            text = self._clean_text('\n'.join(full_text))
            logger.info(f'Extraction [GCS batch]: {filename} ({len(text)} chars)')
            return ExtractionResult(text=text, success=True)
        except Exception as e:
            logger.error(f'Erreur batch GCS {filename}: {e}')
            return ExtractionResult(text='', success=False, error_message=str(e))


    def extract_from_uri(self, gcs_uri: str, filename: str) -> ExtractionResult:
        """Extrait le texte d'un PDF via son URI GCS — sans téléchargement local."""
        if not self._init_document_ai():
            return ExtractionResult(text='', success=False, error_message='Document AI non disponible')
        try:
            from google.cloud import documentai_v1 as documentai
            from google.cloud import storage
            import json, time
            output_prefix = f"_docai_output/{filename.replace('/','_').replace(' ','_')}/"
            output_uri = f"gs://mina-pdfs/{output_prefix}"
            request = documentai.BatchProcessRequest(
                name=self.processor_id,
                input_documents=documentai.BatchDocumentsInputConfig(
                    gcs_documents=documentai.GcsDocuments(
                        documents=[documentai.GcsDocument(gcs_uri=gcs_uri, mime_type='application/pdf')]
                    )
                ),
                document_output_config=documentai.DocumentOutputConfig(
                    gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(gcs_uri=output_uri)
                )
            )
            print(f'[DEBUG] Batch lancé: {gcs_uri}')
            operation = self._docai_client.batch_process_documents(request=request)
            operation.result(timeout=300)
            storage_client = storage.Client()
            bucket = storage_client.bucket('mina-pdfs')
            full_text = []
            for blob in bucket.list_blobs(prefix=output_prefix):
                if blob.name.endswith('.json'):
                    data = json.loads(blob.download_as_text())
                    full_text.append(data.get('text', ''))
                    blob.delete()
            text = self._clean_text('\n'.join(full_text))
            logger.info(f'Extraction [GCS batch]: {filename} ({len(text)} chars)')
            return ExtractionResult(text=text, success=True)
        except Exception as e:
            logger.error(f'Erreur batch GCS {filename}: {e}')
            return ExtractionResult(text='', success=False, error_message=str(e))


    def extract_from_uri(self, gcs_uri: str, filename: str) -> ExtractionResult:
        """Extrait le texte d'un PDF via son URI GCS — sans téléchargement local."""
        if not self._init_document_ai():
            return ExtractionResult(text='', success=False, error_message='Document AI non disponible')
        try:
            from google.cloud import documentai_v1 as documentai
            from google.cloud import storage
            import json, time
            output_prefix = f"_docai_output/{filename.replace('/','_').replace(' ','_')}/"
            output_uri = f"gs://mina-pdfs/{output_prefix}"
            request = documentai.BatchProcessRequest(
                name=self.processor_id,
                input_documents=documentai.BatchDocumentsInputConfig(
                    gcs_documents=documentai.GcsDocuments(
                        documents=[documentai.GcsDocument(gcs_uri=gcs_uri, mime_type='application/pdf')]
                    )
                ),
                document_output_config=documentai.DocumentOutputConfig(
                    gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(gcs_uri=output_uri)
                )
            )
            print(f'[DEBUG] Batch lancé: {gcs_uri}')
            operation = self._docai_client.batch_process_documents(request=request)
            operation.result(timeout=300)
            storage_client = storage.Client()
            bucket = storage_client.bucket('mina-pdfs')
            full_text = []
            for blob in bucket.list_blobs(prefix=output_prefix):
                if blob.name.endswith('.json'):
                    data = json.loads(blob.download_as_text())
                    full_text.append(data.get('text', ''))
                    blob.delete()
            text = self._clean_text('\n'.join(full_text))
            logger.info(f'Extraction [GCS batch]: {filename} ({len(text)} chars)')
            return ExtractionResult(text=text, success=True)
        except Exception as e:
            logger.error(f'Erreur batch GCS {filename}: {e}')
            return ExtractionResult(text='', success=False, error_message=str(e))

    def extract(
        self,
        content: bytes,
        filename: str
    ) -> ExtractionResult:
        """
        Extrait le texte d'un document.
        """
        extension = Path(filename).suffix.lower()
        
        extractors = {
            ".pdf": self._extract_pdf,
            ".txt": self._extract_txt,
            ".docx": self._extract_docx,
            ".doc": self._extract_docx,
        }
        
        extractor = extractors.get(extension)
        
        if extractor is None:
            logger.warning(f"Format non supporté: {extension} ({filename})")
            return ExtractionResult(
                text="",
                success=False,
                error_message=f"Format non supporté: {extension}"
            )
        
        try:
            result = extractor(content, filename)
            logger.info(
                f"Extraction {'[DocAI]' if result.used_ocr else ''}: "
                f"{filename} ({len(result.text)} chars)"
            )
            return result
        except Exception as e:
            logger.error(f"Erreur extraction {filename}: {e}")
            return ExtractionResult(
                text="",
                success=False,
                error_message=str(e)
            )
    
    def _extract_pdf(
        self,
        content: bytes,
        filename: str
    ) -> ExtractionResult:
        """
        Extrait le texte d'un PDF.
        Tente pypdf d'abord, puis Document AI OCR si vide.
        """
        # 1. Essayer pypdf (PDFs texte)
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)
        
        full_text = "\n\n".join(pages_text)
        full_text = self._clean_text(full_text)
        
        # 2. Si pas de texte, utiliser Document AI OCR
        if len(full_text) < 50 and self.use_ocr:
            logger.info(f"PDF scanné détecté, OCR Document AI: {filename}")
            print(f"[DEBUG] OCR Document AI pour: {filename}")
            
            ocr_text = self._extract_with_document_ai(content, filename)
            
            if ocr_text:
                print(f"[DEBUG] OCR succès: {len(ocr_text)} chars")
                return ExtractionResult(
                    text=ocr_text,
                    page_count=len(reader.pages),
                    success=True,
                    used_ocr=True
                )
            else:
                print(f"[DEBUG] OCR échec pour {filename}")
        
        return ExtractionResult(
            text=full_text,
            page_count=len(reader.pages),
            success=True
        )
    
    def _extract_with_document_ai(
        self,
        content: bytes,
        filename: str
    ) -> str:
        """
        Extrait le texte via Google Document AI.
        Traitement côté serveur = pas de problème de mémoire locale.
        """
        if not self._init_document_ai():
            logger.warning("Document AI non disponible")
            return ""
        
        try:
            from google.cloud import documentai_v1 as documentai
            
            # Préparer le document
            gcs_document = documentai.GcsDocument(
                gcs_uri=f"gs://mina-pdfs/{filename}",
                mime_type="application/pdf",
            )
            
            # Créer la requête
            request = documentai.ProcessRequest(
                name=self.processor_id,
                gcs_document=gcs_document,
            )
            
            # Traiter le document (côté serveur Google)
            print(f"[DEBUG] Envoi à Document AI: {len(content)} bytes")
            result = self._docai_client.process_document(request=request)
            
            # Extraire le texte
            document = result.document
            text = document.text
            
            print(f"[DEBUG] Document AI retourne: {len(text)} chars")
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Erreur Document AI: {e}")
            print(f"[DEBUG] Erreur Document AI: {e}")
            return ""
    
    def _extract_txt(
        self,
        content: bytes,
        filename: str
    ) -> ExtractionResult:
        """Extrait le texte d'un fichier TXT."""
        encodings = ["utf-8", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                text = self._clean_text(text)
                return ExtractionResult(text=text, success=True)
            except UnicodeDecodeError:
                continue
        
        text = content.decode("utf-8", errors="ignore")
        return ExtractionResult(
            text=self._clean_text(text),
            success=True,
            error_message="Encodage fallback"
        )
    
    def _extract_docx(
        self,
        content: bytes,
        filename: str
    ) -> ExtractionResult:
        """Extrait le texte d'un fichier DOCX."""
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        
        full_text = "\n\n".join(paragraphs)
        return ExtractionResult(
            text=self._clean_text(full_text),
            success=True
        )
    
    def _clean_text(self, text: str) -> str:
        """Nettoie le texte extrait."""
        import re
        text = re.sub(r" +", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def get_text_extractor(use_ocr: bool = True) -> TextExtractor:
    """Factory pour obtenir un extracteur de texte."""
    return TextExtractor(use_ocr=use_ocr)
